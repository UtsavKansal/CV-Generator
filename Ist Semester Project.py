    
    
from tkinter import *
from tkinter import ttk
import tkinter
import docx
from docx import Document
from PIL import ImageTk, Image
root = Tk()
root.title("Mon CV. Générateur")
frame = Frame(root)
frame.pack(side=TOP)

bottomframe = Frame(root)
bottomframe.pack( side = TOP )
path="logo.png"

img = ImageTk.PhotoImage(Image.open(path))
panel = Label(frame, image = img)
panel.pack(side = TOP)

def mainfunc():
    root.destroy()
    window = Tk()
    window.title("Mon CV. Générateur")
    window.geometry('1000x500')
    
    def itsector():
        cs=Tk()
        cs.title("Mon CV. Générateur")
        cs.geometry('1000x500')
        def itsectorexperienced():
            cs.destroy()
            window1=Tk()
            window1.title("Mon CV. Générateur")
            window1.geometry('1000x500')
            lbl11 = Label(window1, text= 'Name:',anchor=W).grid(column=0, row=0)
            ent11 = Entry(window1,width=20)
            lbl12 = Label(window1, text= "Fathers Name:").grid(column=0, row=1)
            ent12 = Entry(window1,width=20)
            lbl13 = Label(window1, text= "Mother's Name:").grid(column=0, row=2)
            ent13 = Entry(window1,width=20)
            lbl14 = Label(window1, text= "D.O.B.:").grid(column=0, row=3)
            ent14 = Entry(window1,width=20)
            lbl14m = Label(window1, text= '(dd/mm/yy format)').grid(column=2, row=2)
            lbl15 = Label(window1, text= "10th % or C.G.P.A: ").grid(column=0, row=4)
            ent15 = Entry(window1,width=20)
            lbl16 = Label(window1, text= "12th % or C.G.P.A: ").grid(column=0, row=5)
            ent16 = Entry(window1,width=20)
            lbl17 = Label(window1, text= "Hobbies:").grid(column=0, row=6)
            ent17 = Entry(window1,width=20)
            lbl18 = Label(window1, text= 'Academics:').grid(column=0, row=7)
            ent18 = Entry(window1,width=20)
            lbl19 = Label(window1, text= 'C.G.P.A. in all Semesters: ').grid(column=0, row=8)
            ent19 = Entry(window1,width=20)
            lbl20 = Label(window1, text= 'Avg. C.G.P.A.:').grid(column=0, row=9)
            ent20 = Entry(window1,width=20)
            lbl22 = Label(window1, text= 'Years of experience:').grid(column=0, row=11)
            ent22 = Entry(window1,width=20)
            lbl23 = Label(window1, text= 'Knowledge base:').grid(column=0, row=12)
            ent23 = Entry(window1,width=20)
            lbl24 = Label(window1, text= 'Language:').grid(column=0, row=13)
            ent24 = Entry(window1,width=20)
            lbl25= Label(window1, text= 'Other achievement:').grid(column=0, row=14)
            ent25 = Entry(window1,width=20)
            lbl26 = Label(window1, text= 'Technology worked upon:').grid(column=0, row=15)
            ent26 = Entry(window1,width=20)
            lbl27 = Label(window1, text= '1) Languages worked upon:').grid(column=0, row=16)
            ent27 = Entry(window1,width=20)
            lbl28 = Label(window1, text= '2) Operating system:').grid(column=0, row=17)
            ent28 = Entry(window1,width=20)
            lbl29 = Label(window1, text= '3) User interface:').grid(column=0, row=18)
            ent29 = Entry(window1,width=20)
            lbl30 = Label(window1, text= '4) Tools used:').grid(column=0, row=19)
            ent31 = Entry(window1,width=20)
            lbl33 = Label(window1, text= '5) Other technology worked upon:').grid(column=0, row=20)
            ent33 = Entry(window1,width=20)
        
            ent11.grid(column=1, row=0)
            ent12.grid(column=1, row=1)
            ent13.grid(column=1, row=2)
            ent14.grid(column=1, row=3)
            ent15.grid(column=1, row=4)
            ent16.grid(column=1, row=5)
            ent17.grid(column=1, row=6)
            ent18.grid(column=1, row=7)
            ent19.grid(column=1, row=8)
            ent20.grid(column=1, row=9)
            ent22.grid(column=1, row=11)
            ent23.grid(column=1, row=12)
            ent24.grid(column=1, row=13)
            ent25.grid(column=1, row=14)
            ent26.grid(column=1, row=15)
            ent27.grid(column=1, row=16)
            ent28.grid(column=1, row=17)
            ent29.grid(column=1, row=18)
            ent31.grid(column=1, row=19)
            ent33.grid(column=1, row=20)
            
            def cv():
                document=Document()
                document.add_heading('IT JOB RESUME',level=0)
                p=document.add_paragraph()
                p.add_run("Name : ").bold=True
                p.add_run(ent11.get())
                p=document.add_paragraph()
                p.add_run("Father's name : ").bold=True
                p.add_run(ent12.get())
                p=document.add_paragraph()
                p.add_run("Mother's name : ").bold=True
                p.add_run(ent13.get())
                p=document.add_paragraph()
                p.add_run("Date of birth(DOB) : ").bold=True
                p.add_run(ent14.get())
                p=document.add_paragraph()
                p.add_run("Class 10th percentage : ").bold=True
                p.add_run(ent15.get())
                p=document.add_paragraph()
                p.add_run("Class 12th percentage : ").bold=True
                p.add_run(ent16.get())
                p=document.add_paragraph()
                p.add_run("Hobbies : ").bold=True
                p.add_run(ent17.get())
                p=document.add_paragraph()
                p.add_run("Academics : ").bold=True
                p.add_run(ent18.get())
                p=document.add_paragraph()
                #MISSING C.G.P.A in all semesters
                p.add_run("Avg. CGPA : ").bold=True
                p.add_run(ent20.get())
                p=document.add_paragraph()
                p.add_run("Years of experience : ").bold=True
                p.add_run(ent22.get())
                p=document.add_paragraph()
                p.add_run("Knowledge base : ").bold=True
                p.add_run(ent23.get())
                p=document.add_paragraph()
                p.add_run("Language : ").bold=True
                p.add_run(ent24.get())
                p=document.add_paragraph()
                p.add_run("Other achievements : ").bold=True
                p.add_run(ent25.get())
                document.add_heading('TECHNOLOGY WORKED UPON',level=0)
                p=document.add_paragraph()
                p.add_run("Languages worked upon : ").bold=True
                p.add_run(ent27.get())
                p=document.add_paragraph()
                p.add_run("Operating systems worked upon : ").bold=True
                p.add_run(ent28.get())
                p=document.add_paragraph()
                p.add_run("User interface : ").bold=True
                p.add_run(ent29.get())
                p=document.add_paragraph()
                p.add_run("Tools used : ").bold=True
                p.add_run(ent31.get())
                p=document.add_paragraph()
                p.add_run("Other technology worked upon : ").bold=True
                p.add_run(ent33.get())
                document.save('E:/itjob.docx')

            btn1 = tkinter.Button(window1,text = "CLICK HERE TO GENERATE RESUME", fg ="White",bg="black",command=cv).grid(column=0,row=22)
        def itsectorfreshman():
            cs.destroy()
            window1=Tk()
            window1.title("Mon CV. Générateur")
            window1.geometry('500x1100')
            lbl11 = Label(window1, text= 'Name:',anchor=W).grid(column=0, row=0)
            ent11 = Entry(window1,width=20)
            lbl12 = Label(window1, text= 'Address').grid(column=0, row=1)
            ent12 = Entry(window1,width=20)
            lbl13 = Label(window1, text= 'Phone: No.').grid(column=0, row=2)
            ent13 = Entry(window1,width=20)
            lbl14= Label(window1, text= 'Email Address: ').grid(column=0, row=3)
            ent14= Entry(window1,width=20)
            lbl15= Label(window1, text= 'Website(If any)').grid(column=0, row=4)
            ent15= Entry(window1,width=20)
            lbl16= Label(window1, text= 'Father’s name: ').grid(column=0, row=5)
            ent16= Entry(window1,width=20)
            lbl17= Label(window1, text= 'Father’s occupation: ').grid(column=0, row=6)
            ent17= Entry(window1,width=20)
            lbl18= Label(window1, text= 'Mother’s name: ').grid(column=0, row=7)
            ent18= Entry(window1,width=20)
            lbl19= Label(window1, text= 'Mother’s occupation: ').grid(column=0, row=8)
            ent19= Entry(window1,width=20)
            lbl19m   = Label(window1, text="--------------------------------------------").grid(row=9)
            lbl20= Label(window1, text= 'Education: ').grid(column=0, row=10)
            lbl21= Label(window1, text= 'Class 10th marks: ').grid(column=0, row=11)
            ent21= Entry(window1,width=20)
            lbl22= Label(window1, text= 'Class 12th marks: ').grid(column=0, row=12)
            ent22= Entry(window1,width=20)
            lbl23= Label(window1, text= 'Name of degree: ').grid(column=0, row=13)
            ent23= Entry(window1,width=20)
            lbl26= Label(window1, text= 'Avg. C.G.P.A: ').grid(column=0, row=14)
            ent26= Entry(window1,width=20)
            lbl26m = Label(window1, text="--------------------------------------------").grid(row=15)
            lbl27= Label(window1, text= 'Types list of accomplishments: ').grid(column=0, row=16)
            ent28= Entry(window1,width=20)
            ent29= Entry(window1,width=20)
            ent30= Entry(window1,width=20)
            ent31= Entry(window1,width=20)
            lbl31m = Label(window1, text="--------------------------------------------").grid(row=20)
            lbl31= Label(window1, text= 'Internships and project worked upon: ').grid(column=0, row=21)
            lbl31= Label(window1, text= 'Name: ').grid(column=1, row=21)
            lbl31= Label(window1, text= 'Details: ').grid(column=2, row=21)
            ent32= Entry(window1,width=20)
            ent33= Entry(window1,width=20)
            ent34= Entry(window1,width=20)
            ent35= Entry(window1,width=20)
            ent36= Entry(window1,width=20)
            ent37= Entry(window1,width=20)
            lbl37m = Label(window1, text="--------------------------------------------").grid(row=25)
            lbl38= Label(window1, text= 'Type list of skills ').grid(column=0, row=26)
            ent38= Entry(window1,width=20)
            ent39= Entry(window1,width=20)
            ent40= Entry(window1,width=20)
                
            ent11.grid(column=1, row=0)
            ent12.grid(column=1, row=1)
            ent13.grid(column=1, row=2)
            ent14.grid(column=1, row=3)
            ent15.grid(column=1, row=4)
            ent16.grid(column=1, row=5)
            ent17.grid(column=1, row=6)
            ent18.grid(column=1, row=7)
            ent19.grid(column=1, row=8)
            ent21.grid(column=1, row=11)
            ent22.grid(column=1, row=12)
            ent23.grid(column=1, row=13)
            ent26.grid(column=1, row=14)
            ent28.grid(column=1, row=16)
            ent29.grid(column=1, row=17)
            ent30.grid(column=1, row=18)
            ent31.grid(column=1, row=19)
            ent32.grid(column=1, row=22)
            ent33.grid(column=2, row=22)
            ent34.grid(column=1, row=23)
            ent35.grid(column=2, row=23)
            ent36.grid(column=1, row=24)
            ent37.grid(column=2, row=24)
            ent38.grid(column=1, row=26)
            ent39.grid(column=1, row=27)
            ent40.grid(column=1, row=28)
            def cv():
                document=Document()
                document.add_heading('IT RESUME',level=0)
                p=document.add_paragraph()
                p.add_run('Name : ').bold=True
                p.add_run(ent11.get())
                p=document.add_paragraph()
                p.add_run('Address : ').bold=True
                p.add_run(ent12.get())
                p=document.add_paragraph()
                p.add_run('Phone number : ').bold=True
                p.add_run(ent13.get())
                p=document.add_paragraph()
                p.add_run('E-mail id : ').bold=True
                p.add_run(ent14.get())
                p=document.add_paragraph()
                p.add_run('Objectives : ').bold=True
                p.add_run(ent15.get())
                p=document.add_paragraph()
                p.add_run("Father's name : ").bold=True
                p.add_run(ent16.get())
                p=document.add_paragraph()
                p.add_run("Mother's name : ").bold=True
                p.add_run(ent18.get())
                p=document.add_paragraph()
                p.add_run("Father's occupation : ").bold=True
                p.add_run(ent17.get())
                p=document.add_paragraph()
                p.add_run("Mother's occupation : ").bold=True
                p.add_run(ent19.get())
                document.add_heading('Education',level=1)
                p=document.add_paragraph()
                p.add_run('Class 10th percentage : ').bold=True
                p.add_run(ent21.get())
                p=document.add_paragraph()
                p.add_run('Class 12th percentage : ').bold=True
                p.add_run(ent22.get())
                p=document.add_paragraph()
                p.add_run('Name of degree : ').bold=True
                p.add_run(ent23.get())
                p=document.add_paragraph()
                p.add_run('Average CGPA : ').bold=True
                p.add_run(ent26.get())
                p=document.add_paragraph()
                p.add_run('Accomplishments : ').bold=True
                p.add_run(ent28.get()+ent29.get()+ent30.get()+ent31.get())
                document.add_heading('Internships and Projects worked upon',level=1)
                table=document.add_table(rows=4,cols=2)
                row=table.rows[0]
                row.cells[0].text=('NAME')
                row.cells[1].text=('DETAILS')
                row=table.rows[1]
                row.cells[0].text=ent32.get()
                row.cells[1].text=ent33.get()
                row=table.rows[2]
                row.cells[0].text=ent34.get()
                row.cells[1].text=ent35.get()
                row=table.rows[3]
                row.cells[0].text=ent36.get()
                row.cells[1].text=ent37.get()
                document.add_heading('Skills',level=1)
                p=document.add_paragraph()
                p.add_run(ent38.get()+ent39.get()+ent40.get())
                document.save('E:/itfreshman.docx')
            btn1 = tkinter.Button(window1,text = "CLICK HERE TO GENERATE RESUME", fg ="White",bg="black",command=cv).grid(column=0,row=32)
        
        btn1 = tkinter.Button(cs,text = "FRESHMAN", fg ="White",bg="black",command=itsectorfreshman)
        btn2 = tkinter.Button(cs,text ="EXPERIENCED", fg ="White",bg="black" ,command=itsectorexperienced)
        btn1.pack(side="top")
        btn2.pack(side="top")
        btn1.config(height=5,width=50)
        btn2.config(height=5,width=50)
        
    
    def ssbinterviews():
        
        window1=Tk()
        window1.geometry('500x1000')
        window1.title("Mon CV. Générateur")
        lbl11 = Label(window1, text= 'Name:').grid(column=0, row=0)
        ent11 = Entry(window1,width=20)
        lbl12 = Label(window1, text= "Date of birth:").grid(column=0, row=1)
        ent12 = Entry(window1,width=20)
        lbl13 = Label(window1, text= "Age:").grid(column=0, row=2)
        ent13 = Entry(window1,width=20)
        lbl13 = Label(window1, text= "Height:").grid(column=0, row=3)
        ent14 = Entry(window1,width=20)
        lbl13 = Label(window1, text= "Weight: ").grid(column=0, row=4)
        ent15 = Entry(window1,width=20)
        lbl13 = Label(window1, text= "Marital status: ").grid(column=0, row=5)
        ent16 = Entry(window1,width=20)
        lbl13 = Label(window1, text= "Religion:").grid(column=0, row=6)
        ent100 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Mother tongue:').grid(column=0, row=7)
        ent17 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Community: ').grid(column=0, row=8)
        ent101 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Place of maximum residence (with residence):').grid(column=0, row=9)
        ent18 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Projects worked upon in companies:').grid(column=0, row=10)
        ent19 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Present residence:').grid(column=0, row=11)
        ent20 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Permanent address:').grid(column=0, row=12)
        ent21 = Entry(window1,width=20)
        lbl14= Label(window1, text= 'Father’s occupation:').grid(column=0, row=13)
        ent22 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Mother’s occupation:').grid(column=0, row=14)
        ent23 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Father’s income:').grid(column=0, row=15)
        ent24 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Mother’s income:').grid(column=0, row=16)
        ent25 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Graduation:').grid(column=0, row=17)
        ent26 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Educational qualifications:').grid(column=0, row=18)
        ent27 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Phone number:').grid(column=0, row=19)
        ent28 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Mobile number:').grid(column=0, row=20)
        ent29 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Hobbies:').grid(column=0, row=21)
        ent30 = Entry(window1,width=20)
        lbl   = Label(window1, text="--------------------------------------------").grid(row=23)
        lbl14 = Label(window1, text= 'NCC’s training:').grid(column=0, row=24)
        ent31 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Wing:').grid(column=0, row=25)
        ent102 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'Achievements if any:').grid(column=0, row=26)
        ent32 = Entry(window1,width=20)
        
        
        ent11.grid(column=1, row=0)
        ent12.grid(column=1, row=1)
        ent13.grid(column=1, row=2)
        ent14.grid(column=1, row=3)
        ent15.grid(column=1, row=4)
        ent16.grid(column=1, row=5)
        ent100.grid(column=1, row=6)
        ent17.grid(column=1, row=7)
        ent101.grid(column=1, row=8)
        ent18.grid(column=1, row=9)
        ent19.grid(column=1, row=10)
        ent20.grid(column=1, row=11)
        ent21.grid(column=1, row=12)
        ent22.grid(column=1, row=13)
        ent23.grid(column=1, row=14)
        ent24.grid(column=1, row=15)
        ent25.grid(column=1, row=16)
        ent26.grid(column=1, row=17)
        ent27.grid(column=1, row=18)
        ent28.grid(column=1, row=19)
        ent29.grid(column=1, row=20)
        ent30.grid(column=1, row=21)
        ent31.grid(column=1, row=24)
        ent102.grid(column=1,row=25)
        ent32.grid(column=1, row=26)
        
        def cv():
            document=Document()
            document.add_heading('SSB RESUME',level=0)
            p=document.add_paragraph()
            p.add_run('Name : ').bold=True
            p.add_run(ent11.get())
            p=document.add_paragraph()
            p.add_run('Date of birth : ').bold=True
            p.add_run(ent12.get())
            p=document.add_paragraph()
            p.add_run('Age : ').bold=True
            p.add_run(ent13.get())
            p=document.add_paragraph()
            p.add_run('Height : ').bold=True
            p.add_run(ent14.get())
            
            p=document.add_paragraph()
            p.add_run('Weight : ').bold=True
            p.add_run(ent15.get())
            p=document.add_paragraph()
            p.add_run('Marital status : ').bold=True
            p.add_run(ent16.get())
            p=document.add_paragraph()
            p.add_run("Religion : ").bold=True
            p.add_run(ent100.get())
            p=document.add_paragraph()
            p.add_run("Mother tongue : ").bold=True
            p.add_run(ent17.get())
            p=document.add_paragraph()
            p.add_run('Community : ').bold=True
            p.add_run(ent101.get())
            p=document.add_paragraph()
            p.add_run('Place of maximum residence : ').bold=True
            p.add_run(ent18.get())
            p=document.add_paragraph()
            #MISSING project worked upon
            p.add_run('Present residence : ').bold=True
            p.add_run(ent20.get())
            p=document.add_paragraph()
            #MISSING permanent address
            p.add_run('Father’s occupation : ').bold=True
            p.add_run(ent22.get())
            p=document.add_paragraph()
            p.add_run('Mother’s occupation : ').bold=True
            p.add_run(ent23.get())
            p=document.add_paragraph()
            p.add_run('Father’s income : ').bold=True
            p.add_run(ent24.get())
            p=document.add_paragraph()
            p.add_run('Mother’s income : ').bold=True
            p.add_run(ent25.get())
            p=document.add_paragraph()
            p.add_run('Graduation : ').bold=True
            p.add_run(ent26.get())
            p=document.add_paragraph()
            p.add_run('Educational qualifications : ').bold=True
            p.add_run(ent27.get())
            p=document.add_paragraph()
            p.add_run('Phone number : ').bold=True
            p.add_run(ent28.get())
            p=document.add_paragraph()
            p.add_run('Mobile number : ').bold=True
            p.add_run(ent29.get())
            p=document.add_paragraph()
            p.add_run('Hobbies : ').bold=True
            p.add_run(ent30.get())
            p=document.add_paragraph()
            p=document.add_paragraph()
            p.add_run('NCC’s training : ').bold=True
            p.add_run(ent31.get())
            p=document.add_paragraph()
            p.add_run('Wing : ').bold=True
            p.add_run()
            p=document.add_paragraph()
            p.add_run('Achievements (if any) : ').bold=True
            p.add_run(ent32.get())
            document.save('E:/ssb.docx')
        
        btn1 = tkinter.Button(window1,text = "CLICK HERE TO GENERATE RESUME", fg ="White",bg="black",command=cv).grid(column=0,row=27)
    def civilservices():
        cs=Tk()
        cs.geometry('1000x500')
        cs.title("Mon CV. Générateur")
        def civilservicesExp():
            cs.destroy()
            window1=Tk()
            window1.title("Mon CV. Générateur")
            window1.geometry('1000x500')
            lbl11 = Label(window1, text='Name:').grid(column=0, row=0)
            ent11 = Entry(window1, width=20)
            lbl12 = Label(window1, text="Meaning of name:").grid(column=0, row=1)
            ent12 = Entry(window1, width=20)
            lbl13 = Label(window1, text="D.O.B:").grid(column=0, row=2)
            ent13 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Community:").grid(column=0, row=3)
            ent14 = Entry(window1, width=20)
            lbl13m = Label(window1, text='(dd/mm/yy format)',anchor=W).grid(column=2, row=2,)
            lbl13 = Label(window1, text="Nationality: ").grid(column=0, row=4)
            ent15 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Gender: ").grid(column=0, row=5)
            ent16 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Objectives:").grid(column=0, row=6)
            ent17 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Educational Qualifications:').grid(column=0, row=7)
            ent18 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Hobbies: ').grid(column=0, row=8)
            ent19 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Present designation(if any):').grid(column=0, row=9)
            ent20 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Telephone no.:').grid(column=0, row=10)
            ent21 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Mobile no.:').grid(column=0, row=11)
            ent22 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Please indicate whether serving under: ').grid(column=0, row=12)
            ent23 = Entry(window1, width=20)
            lbl11111=Label(window1, text='central government/ state govt./autonomous organisation / university / any other institution / private Organisation: ').grid(column=2, row=12)
            lbl14 = Label(window1, text='Present designation(if serving):').grid(column=0, row=13)
            ent24 = Entry(window1, width=20)
    
    
            
            ent11.grid(column=1, row=0)
            ent12.grid(column=1, row=1)
            ent13.grid(column=1, row=2)
            ent14.grid(column=1, row=3)
            ent15.grid(column=1, row=4)
            ent16.grid(column=1, row=5)
            ent17.grid(column=1, row=6)
            ent18.grid(column=1, row=7)
            ent19.grid(column=1, row=8)
            ent20.grid(column=1, row=9)
            ent21.grid(column=1, row=10)
            ent22.grid(column=1, row=11)
            ent23.grid(column=1, row=12)
            ent24.grid(column=1,row=13)
            
            def cv():
                document=Document()
                document.add_heading('CIVIL SERVICES RESUME',level=0)
                p=document.add_paragraph()
                p.add_run('Name : ').bold=True
                p.add_run(ent11.get())
                p=document.add_paragraph()
                p.add_run("Meaning of name : ").bold=True
                p.add_run(ent12.get())
                p=document.add_paragraph()
                p.add_run("Date of birth : ").bold=True
                p.add_run(ent13.get())
                p=document.add_paragraph()
                p.add_run("Community : ").bold=True
                p.add_run(ent14.get())
                p=document.add_paragraph()
                p.add_run("Nationality : ").bold=True
                p.add_run(ent15.get())
                p=document.add_paragraph()
                p.add_run("Gender : ").bold=True
                p.add_run(ent16.get())
                p=document.add_paragraph()
                p.add_run("Objectives : ").bold=True
                p.add_run(ent17.get())
                p=document.add_paragraph()
                p.add_run("Educational qualifications : ").bold=True
                p.add_run(ent18.get())
                p=document.add_paragraph()
                p.add_run("Hobbies : ").bold=True
                p.add_run(ent19.get())
                p=document.add_paragraph()
                p.add_run("Present designation (if any) : ").bold=True
                p.add_run(ent20.get())
                p=document.add_paragraph()
                p.add_run("Telephone number : ").bold=True
                p.add_run(ent21.get())
                p=document.add_paragraph()
                p.add_run("Mobile number : ").bold=True
                p.add_run(ent22.get())
                p=document.add_paragraph()
                p.add_run("Please indicate whether serving under central government / state govt. / autonomous organisation / university / any other institution / private Organisation / central or state government undertaking or self-employed.").bold=True
                p=document.add_paragraph(ent23.get())
                p=document.add_paragraph()
                p.add_run("Present designation (if serving) : ").bold=True
                p.add_run(ent24.get())
                document.save('E:/civilexperience.docx')
            
            btn1 = tkinter.Button(window1,text = "CLICK HERE TO GENERATE RESUME", fg ="White",bg="black",command=cv).grid(column=0,row=22)
            
        def civilservicesWithoutexp():
            cs.destroy()
            window1 = Tk()
            window1.geometry('1000x500')
            window1.title("Mon CV. Générateur")
            lbl11 = Label(window1, text='Name:').grid(column=0, row=0)
            ent11 = Entry(window1, width=20)
            lbl12 = Label(window1, text="Fathers Name:").grid(column=0, row=1)
            ent12 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Mother's Name:").grid(column=0, row=2)
            ent13 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Meaning of name:").grid(column=0, row=3)
            ent14 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Father's occupation :").grid(column=0, row=4)
            ent15 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Mother's occupation:").grid(column=0, row=5)
            ent16 = Entry(window1, width=20)
            lbl13m = Label(window1, text='(dd/mm/yy format)').grid(column=2, row=8)
            lbl13 = Label(window1, text="Educational Qualification: ").grid(column=0, row=6)
            ent17 = Entry(window1, width=20)
            lbl13 = Label(window1, text="Hobbies: ").grid(column=0, row=7)
            ent18 = Entry(window1, width=20)
            lbl13 = Label(window1, text="D.O.B:").grid(column=0, row=8)
            ent19 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Language proficiency:').grid(column=0, row=9)
            ent20 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Community: ').grid(column=0, row=10)
            ent21 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Nationality:').grid(column=0, row=11)
            ent22 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Gender:').grid(column=0, row=12)
            ent23 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Class 12th %:').grid(column=0, row=13)
            ent24 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Class 10th %:').grid(column=0, row=14)
            ent25 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Extracurricular activites:').grid(column=0, row=15)
            ent26 = Entry(window1, width=20)
            lbl14 = Label(window1, text='certificate of interest:').grid(column=0, row=16)
            ent27 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Degree earned:').grid(column=0, row=17)
            ent28 = Entry(window1, width=20)
            lbl14 = Label(window1, text='Objectives:').grid(column=0, row=18)
            ent29 = Entry(window1, width=20)
            
            ent11.grid(column=1, row=0)
            ent12.grid(column=1, row=1)
            ent13.grid(column=1, row=2)
            ent14.grid(column=1, row=3)
            ent15.grid(column=1, row=4)
            ent16.grid(column=1, row=5)
            ent17.grid(column=1, row=6)
            ent18.grid(column=1, row=7)
            ent19.grid(column=1, row=8)
            ent20.grid(column=1, row=9)
            ent21.grid(column=1, row=10)
            ent22.grid(column=1, row=11)
            ent23.grid(column=1, row=12)
            ent24.grid(column=1, row=13)
            ent25.grid(column=1, row=14)
            ent26.grid(column=1, row=15)
            ent27.grid(column=1, row=16)
            ent28.grid(column=1, row=17)
            ent29.grid(column=1, row=18)
            
            def cv():
                document=Document()
                document.add_heading('CIVIL SERVICES EXAM RESUME',level=0)
                p=document.add_paragraph()
                p.add_run('Name : ').bold=True
                p.add_run(ent11.get())
                p=document.add_paragraph()
                p.add_run("Meaning of name : ").bold=True
                p.add_run(ent14.get())
                p=document.add_paragraph()
                p.add_run("Father’s name : ").bold=True
                p.add_run(ent12.get())
                p=document.add_paragraph()
                p.add_run("Mother’ name : ").bold=True
                p.add_run(ent13.get())
                p=document.add_paragraph()
                p.add_run("Father’s occupation : ").bold=True
                p.add_run(ent15.get())
                p=document.add_paragraph()
                p.add_run("Mother’s occupation : ").bold=True
                p.add_run(ent16.get())
                p=document.add_paragraph()
                p.add_run("Educational qualifications : ").bold=True
                p.add_run(ent17.get())
                p=document.add_paragraph()
                p.add_run("Hobbies : ").bold=True
                p.add_run(ent18.get())
                p=document.add_paragraph()
                p.add_run('Date of birth : ').bold=True
                p.add_run(ent19.get())
                p=document.add_paragraph()
                p.add_run('Language proficiency : ').bold=True
                p.add_run(ent20.get())
                p=document.add_paragraph()
                p.add_run('Community : ').bold=True
                p.add_run(ent21.get())
                p=document.add_paragraph()
                p.add_run('Nationality : ').bold=True
                p.add_run(ent22.get())
                p=document.add_paragraph()
                p.add_run('Gender : ').bold=True
                p.add_run(ent23.get())
                p=document.add_paragraph()
                p.add_run('Class 12th percentage : ').bold=True
                p.add_run(ent24.get())
                p=document.add_paragraph()
                p.add_run('Class 12th percentage : ').bold=True
                p.add_run(ent25.get())
                p=document.add_paragraph()
                p.add_run('Extracurricular activities : ').bold=True
                p.add_run(ent26.get())
                p=document.add_paragraph()
                p.add_run('Certificate of interest : ').bold=True
                p.add_run(ent27.get())
                p=document.add_paragraph()
                p.add_run('Degree earned : ').bold=True
                p.add_run(ent28.get())
                p=document.add_paragraph()
                p.add_run('Objectives : ').bold=True
                p.add_run(ent29.get())
                document.save('E:/civilexam.docx')

            btn1 = tkinter.Button(window1,text = "CLICK HERE TO GENERATE RESUME", fg ="White",bg="black",command=cv).grid(column=0,row=19)
            
        btn1 = tkinter.Button(cs,text = "EXPERIENCE", fg ="White",bg="black",command=civilservicesExp)
        btn2 = tkinter.Button(cs,text ="WITHOUT EXPERIENCE", fg ="White",bg="black" ,command=civilservicesWithoutexp)
        btn1.pack(side=TOP)
        btn2.pack(side=TOP)
        btn1.config(height=5,width=50)
        btn2.config(height=5,width=50)
    def biodata():
        window1=Tk()
        window1.geometry('1000x500')
        window1.title("Mon CV. Générateur")
        lbl11 = Label(window1, text= 'NAME:',anchor=W).grid(column=0, row=0)
        ent11 = Entry(window1,width=20)
        lbl12 = Label(window1, text= "FATHER's NAME:").grid(column=0, row=1)
        ent12 = Entry(window1,width=20)
        lbl13 = Label(window1, text= 'MOTHER’S NAME:').grid(column=0, row=2)
        ent13 = Entry(window1,width=20)
        lbl14 = Label(window1, text= 'FATHER’S OCCUPATION:').grid(column=0, row=3)
        ent14 = Entry(window1,width=20)
        lbl15 = Label(window1, text= 'MOTHER’S OCCUPATION:').grid(column=0, row=4)
        ent15 = Entry(window1,width=20)
        lbl16 = Label(window1, text= 'AGE:').grid(column=0, row=5)
        ent16 = Entry(window1,width=20)
        lbl18 = Label(window1, text= 'HEIGHT:').grid(column=0, row=6)
        ent18 = Entry(window1,width=20)
        lbl19 = Label(window1, text= 'CLASS 10th MARKS:').grid(column=0, row=7)
        ent19 = Entry(window1,width=20)
        lbl110= Label(window1, text= 'CLASS 12th MARKS:').grid(column=0, row=8)
        ent110= Entry(window1,width=20)
        lbl111= Label(window1, text= 'EDUCATIONAL QUALIFICATIONS:').grid(column=0, row=9)
        ent111= Entry(window1,width=20)
        lbl112= Label(window1, text= 'HOBBIES:').grid(column=0, row=10)
        ent112 = Entry(window1,width=20)
        lbl113 = Label(window1, text= 'OBJECTIVE OF LIFE:').grid(column=0, row=11)
        ent113 = Entry(window1,width=20)
        ent11.grid(column=1, row=0)
        ent12.grid(column=1, row=1)
        ent13.grid(column=1, row=2)
        ent14.grid(column=1, row=3)
        ent15.grid(column=1, row=4)
        ent16.grid(column=1, row=5)
        ent18.grid(column=1, row=6)
        ent19.grid(column=1, row=7)
        ent110.grid(column=1, row=8)
        ent111.grid(column=1, row=9)
        ent112.grid(column=1, row=10)
        ent113.grid(column=1, row=11)

        def cv():
            document=Document()
            document.add_heading('RESUME',level=0)
            p=document.add_paragraph()
            p.add_run('Name : ').bold=True
            p.add_run(ent11.get())
            p=document.add_paragraph()
            p.add_run("Father's Name : ").bold=True
            p.add_run(ent12.get())
            p=document.add_paragraph()
            p.add_run("Mother's Name : ").bold=True
            
            p.add_run(ent13.get())
            p=document.add_paragraph()
            p.add_run("Father's occupation : ").bold=True
            p.add_run(ent14.get())
            p=document.add_paragraph()
            p.add_run("Mother's occupation : ").bold=True
            p.add_run(ent15.get())
            p=document.add_paragraph()
            p.add_run("Age : ").bold=True
            p.add_run(ent16.get())
            p=document.add_paragraph()
            p.add_run("Height : ").bold=True
            p.add_run(ent18.get())
            p=document.add_paragraph()
            p.add_run("Class 10th Percentage : ").bold=True
            p.add_run(ent19.get())
            p=document.add_paragraph()
            p.add_run("Class 12th Percentage : ").bold=True
            p.add_run(ent110.get())
            p=document.add_paragraph()
            p.add_run("Educational qualifications : ").bold=True
            p.add_run(ent111.get())
            p=document.add_paragraph()
            p.add_run("Hobbies : ").bold=True
            p.add_run(ent112.get())
            p=document.add_paragraph()
            p.add_run("Objective of life : ").bold=True
            p.add_run(ent113.get())
            document.save('E:/simplebiodata.docx')
        btn1 = tkinter.Button(window1,text = "CLICK HERE TO GENERATE RESUME", fg ="White",bg="black",command=cv).grid(column=0,row=12)
        
        
    btn1 = tkinter.Button(window, text = "SSB", fg ="White",bg="black",command=ssbinterviews)
    btn2 = tkinter.Button(window, text ="I.T SECTOR", fg ="White",bg="black" ,command=itsector)
    btn3 = tkinter.Button(window, text ="BIO DATA", fg = "White", bg="black",command=biodata)
    btn4 = tkinter.Button(window, text = "CIVIL SERVICES", fg = "White", bg="black",command=civilservices)
    btn1.pack(side=TOP)
    btn2.pack(side=TOP)
    btn3.pack(side=TOP)
    btn4.pack(side=TOP)
    btn1.config(height=5,width=50)
    btn2.config(height=5,width=50)
    btn3.config(height=5,width=50)
    btn4.config(height=5,width=50)
    
    window.mainloop()
redbutton = Button(bottomframe, text="CLICK HERE TO BEGIN",bg="yellow",fg="red",command=mainfunc)
redbutton.pack(side=TOP)
redbutton.config( height = 5, width = 50)

root.mainloop()